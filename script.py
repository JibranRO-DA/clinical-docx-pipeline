"""
ANVA - Clinical .docx Processing Pipeline

Extracts structured information from medical notes in .docx format
for epidemiological analysis and reporting.
"""

from docx import Document
import re
import zipfile
import xml.etree.ElementTree as ET
import pandas as pd


def infoall(nm):
    try:
        doc=Document(nm)
    except (zipfile.BadZipFile, ValueError, KeyError, Exception):
        listavacia=["Nombre",'Servicio','Fecha','Edad','Sexo','Diagnostico']
        anva_dfx=pd.DataFrame(columns=listavacia)
        return anva_dfx
        
        
    docomp = "\n".join([p.text for p in doc.paragraphs])
    
    Servicio = doc.paragraphs[0].text.strip()
    it=[0,1,2]
    for i in it:
        Servicio = doc.paragraphs[i].text.strip()
        if Servicio:
            break
        
        
    
    if Servicio == 'NOTA MÉDICA' or Servicio == 'EXAMEN MEDICO' or Servicio == 'EXAMEN MÉDICO':
        
        Fecha1 = re.search(r'(\b\d{1,2}\s*[\/\.]\s*\d{1,2}\s*[\/\.]\s*\d{2,4}\b)',docomp)
        
        if Fecha1 == None:
            Fecha = 'No encontrada'
        else:
            Fecha = Fecha1.group(1)
            
        Apaterno = re.search(r"Apellido paterno:\s*(\w+)", docomp)
        Amaterno = re.search(r"Apellido materno:\s*(\w+)", docomp)
        Nombres = re.search(r"Nombre\(s\)\s*:\s*(.+)", docomp)
        
        if Nombres and Amaterno and Apaterno:
            Nomcomp = f"{Nombres.group(1).strip()} {Apaterno.group(1)} {Amaterno.group(1)}"
        else:
            Nomcomp="No encontrado"
        #Nomcomp = f"{Nombres.group(1).strip()} {Apaterno.group(1)} {Amaterno.group(1)}" if Nombres else "No encontrado"
        
        
        edad1 = re.search(r"Edad:\s*(\d+)", docomp)
        if edad1 == None:
            edad='No encontrada'
        else:
            edad=edad1.group(1).strip()
        
        
        # para sexo
        
        docx_path = nm

        namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "w14": "http://schemas.microsoft.com/office/word/2010/wordml"}

        with zipfile.ZipFile(docx_path) as z:
            xml_content = z.read("word/document.xml")

        root = ET.fromstring(xml_content)

        for para in root.findall(".//w:p", namespaces):

            texts = [t.text for t in para.findall(".//w:t", namespaces) if t.text]
            full_text = " ".join(texts)

            # Solo procesar la línea que contiene "Sexo"
            if "Sexo" not in full_text:
                continue

            checkboxes = []

            for cb in para.findall(".//w14:checkbox", namespaces):
                checked = cb.find(".//w14:checked", namespaces)

                if checked is not None and checked.attrib.get("{http://schemas.microsoft.com/office/word/2010/wordml}val") == "1":
                    checkboxes.append("\u2612")  # ☒
                else:
                    checkboxes.append("\u2610")  # ☐

            sex = full_text, " ".join(checkboxes)
            sex = str(sex)
        

            sexbus = str(re.findall(r"\w+\s+☒",sex))[2:12].strip()
            
        paciente={"Nombre": Nomcomp.upper(),'Servicio': Servicio,'Fecha':Fecha,'Edad': edad,'Sexo': sexbus}
    
        #para diagnostico
        
        
        Dx=[]
        if Servicio == 'NOTA MÉDICA':
            for p in doc.paragraphs[20:30]:
                texto=p.text.strip()
                palas=texto.split()
                np=len(palas)
                if 1 <= np <= 4 and texto.isupper():
                    Dx.append(p.text)

        else:
            Dx.append('EXAMEN MÉDICO')

        registro=[]
        for d in Dx:
            nreg=paciente.copy()
            nreg['Diagnostico']=d
            registro.append(nreg)
            
        anva_dfx=pd.DataFrame(registro)    
        
    else:
        listavacia=["Nombre",'Servicio','Fecha','Edad','Sexo','Diagnostico']
        anva_dfx=pd.DataFrame(columns=listavacia)
            
            
            
        

                    
    return anva_dfx


# =========================
# Main execution
# =========================

import os
base_path = os.getcwd()
os.chdir(base_path)
print(os.getcwd())




lpa=os.listdir()
dfconcat=[]
for pa in lpa:
    os.chdir(pa)
    ldoc=os.listdir()
    for do in ldoc:
        if do.endswith('.docx'):
            print(f"procesando: {do}")
            df_extra=infoall(do)
            dfconcat.append(df_extra)
    os.chdir('..')
    print(pa,' Terminado...')
    
df_fin=pd.concat(dfconcat, ignore_index=True)

os.getcwd()

df_fin.to_excel('Compilado pacientes.xlsx',index=False)

